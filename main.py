"""
Orchestratore principale: lezioni universitarie → dispense Word.
Chiede il file di partenza e coordina trascrizione, elaborazione LLM e generazione Word.
"""

import os
import sys
from pathlib import Path

# Carica variabili da .env o gemini.env (GEMINI_API_KEY) prima di importare gli altri moduli
from dotenv import load_dotenv
_project_dir = Path(__file__).resolve().parent
load_dotenv(_project_dir / ".env")
load_dotenv(_project_dir / "gemini.env")  # supporto anche gemini.env

from transcriber import transcribe_media
from gemini_engine import process_with_gemini
from doc_builder import build_document


SUPPORTED_EXTENSIONS = {".mp4", ".mkv", ".mp3", ".wav", ".m4a", ".txt", ".pdf", ".docx"}


def get_input_file() -> Path:
    """Chiede all'utente il percorso del file di partenza (video, audio, testo o PDF)."""
    print("=" * 60)
    print("  Lezioni → Dispense Word")
    print("=" * 60)
    print("\nFormati supportati: video (.mp4, .mkv), audio (.mp3, .wav, .m4a), testo (.txt), PDF (.pdf), Word (.docx)\n")

    while True:
        path_str = input("Inserisci il percorso del file di partenza: ").strip().strip('"')
        if not path_str:
            print("Percorso vuoto. Riprova.")
            continue

        path = Path(path_str)
        if not path.exists():
            print(f"File non trovato: {path}")
            continue

        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            print(f"Estensione non supportata: {path.suffix}. Usa: {', '.join(SUPPORTED_EXTENSIONS)}")
            continue

        return path.resolve()


def get_text_from_file(path: Path) -> str:
    """Legge il testo da file .txt, .pdf o .docx (senza trascrizione)."""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            print("Per i PDF installa: pip install pypdf")
            sys.exit(1)
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            print("Per i file Word installa: pip install python-docx")
            sys.exit(1)

        doc = Document(str(path))

        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        # Include anche testo in tabelle (alcune esportazioni lo usano)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)

        return "\n".join(parts)
    return ""


def run_workflow(input_path: Path) -> None:
    """Esegue il flusso: trascrizione (se media) → Gemini → Word."""
    transcript_path = input_path.with_name(f"{input_path.stem}_trascrizione.txt")
    output_path = input_path.with_name(f"{input_path.stem}_dispensa.docx")
    transcript_created = False

    print("\n" + "=" * 50)
    print("  Step 1/4: Analisi file di partenza")
    print("=" * 50)
    print(f"  File: {input_path.name}")

    # Testo: da file; Media: da trascrizione
    if input_path.suffix.lower() in {".txt", ".pdf", ".docx"}:
        print("  Lettura testo da file...")
        raw_text = get_text_from_file(input_path)
        if not raw_text.strip():
            print("  Errore: file vuoto o testo non estraibile.")
            return
        print(f"  Testo estratto ({len(raw_text)} caratteri).")
    else:
        print("  Trascrizione audio/video in corso...", flush=True)
        sys.stdout.flush()
        sys.stderr.flush()
        try:
            # Passiamo il path: la trascrizione viene salvata dentro transcribe_media prima del return
            raw_text = transcribe_media(input_path, transcript_save_path=transcript_path)
        except Exception as e:
            print(f"  Errore durante la trascrizione: {e}", flush=True)
            raise
        sys.stdout.flush()
        sys.stderr.flush()
        # Fallback: a volte il trascrittore salva su file ma il valore ritornato può risultare vuoto.
        if not raw_text or not raw_text.strip():
            if transcript_path.exists():
                try:
                    raw_text = transcript_path.read_text(encoding="utf-8", errors="replace")
                    print("  Trascrizione letta dal file di output (fallback).", flush=True)
                except Exception as e:
                    print(f"  Attenzione: lettura trascrizione da file fallita: {e}", flush=True)
        if not raw_text or not raw_text.strip():
            print("  Errore: trascrizione vuota. Verifica il file.", flush=True)
            return
        print(f"  Trascrizione completata ({len(raw_text)} caratteri).", flush=True)
        # Se il trascriber ha già salvato, non riscrivere
        if not transcript_path.exists():
            try:
                transcript_path.write_text(raw_text, encoding="utf-8")
                print(f"  Trascrizione salvata: {transcript_path}", flush=True)
            except Exception as e:
                print(f"  Attenzione: impossibile salvare la trascrizione su file: {e}", flush=True)
        transcript_created = transcript_path.exists()
        sys.stdout.flush()

    print("\n" + "=" * 50, flush=True)
    print("  Step 2/4: Elaborazione con Gemini")
    print("=" * 50)
    print("  Invio testo a Gemini per strutturazione e riordino...", flush=True)
    sys.stdout.flush()
    try:
        structured_content = process_with_gemini(raw_text)
    except Exception as e:
        print(f"  Errore Gemini: {e}")
        print("  La trascrizione è stata salvata; puoi riprovare dopo o usare il file .txt.")
        structured_content = None

    if not structured_content:
        # Fallback: crea comunque un .docx anche se Gemini fallisce/parsa male.
        print("  Avviso: Gemini non ha restituito una struttura valida; genero comunque un Word con contenuto grezzo.", flush=True)
        structured_content = {
            "title": "Dispensa",
            "introduction": "",
            "toc": ["Contenuto"],
            "chapters": [{"title": "Contenuto", "content": raw_text}],
        }
    print("  Risposta ricevuta. Struttura pronta per il documento.")

    print("\n" + "=" * 50)
    print("  Step 3/4: Generazione documento Word")
    print("=" * 50)
    print("  Creazione dispensa (titolo, introduzione, indice, capitoli)...")
    try:
        build_document(structured_content, output_path)
        print(f"  Dispensa salvata: {output_path}")
    except Exception as e:
        print(f"  Errore creazione Word: {e}")
        raise

    print("\n" + "=" * 50)
    print("  Step 4/4: Completato")
    print("=" * 50)
    if transcript_created:
        print(f"  File creati:")
        print(f"    - Trascrizione: {transcript_path}")
    print(f"    - Dispensa Word: {output_path}")
    print()


def main() -> None:
    input_path = get_input_file()
    try:
        run_workflow(input_path)
    except Exception as e:
        print("\n  ! Errore durante l'esecuzione:", flush=True)
        print(f"  {e}", flush=True)
        import traceback
        traceback.print_exc()
        sys.stdout.flush()
        sys.stderr.flush()
        sys.exit(1)


if __name__ == "__main__":
    main()
