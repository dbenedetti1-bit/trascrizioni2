"""
Costruttore del documento Word: crea la dispensa con struttura accademica
(Titolo, Introduzione, Indice, Capitoli) usando python-docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_document(structured_content: dict, output_path: Path) -> None:
    """
    Genera un file .docx a partire dal dizionario restituito da Gemini.
    Struttura: Titolo, Introduzione, Indice dei contenuti, Capitoli sviluppati.
    """
    doc = Document()
    output_path = Path(output_path)
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    title = structured_content.get("title") or "Dispensa"
    introduction = structured_content.get("introduction") or ""
    toc = structured_content.get("toc") or []
    chapters = structured_content.get("chapters") or []

    # Stile titolo principale
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduzione
    doc.add_heading("Introduzione", level=1)
    doc.add_paragraph(introduction)

    # Indice dei contenuti
    doc.add_heading("Indice dei contenuti", level=1)
    for i, entry in enumerate(toc, start=1):
        doc.add_paragraph(f"{i}. {entry}", style="List Number")

    # Capitoli
    for ch in chapters:
        ch_title = ch.get("title") or "Capitolo"
        ch_content = ch.get("content") or ""
        doc.add_heading(ch_title, level=1)
        doc.add_paragraph(ch_content)

    doc.save(str(output_path))
